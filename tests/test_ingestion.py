"""Tests for document ingestion."""

from __future__ import annotations

import hashlib
from pathlib import Path

import pytest
from docx import Document

from app.ingestion.extractor import compute_checksum_sha256, extract_document


def test_txt_extraction(tmp_path: Path) -> None:
    sample = tmp_path / "notes.txt"
    content = "TraceDoc ingestion test.\nSecond line."
    sample.write_text(content, encoding="utf-8")

    result = extract_document(str(sample))

    assert result.text == content
    assert result.file_type == "txt"
    assert result.file_name == "notes.txt"
    assert result.page_count is None
    assert result.metadata.get("encoding") == "utf-8"
    assert result.extraction_warnings == []


def test_txt_latin1_fallback_warning(tmp_path: Path) -> None:
    sample = tmp_path / "latin.txt"
    sample.write_bytes(b"\xff\xfe invalid utf-8 but latin-1 ok")

    result = extract_document(str(sample))

    assert result.metadata.get("encoding") == "latin-1"
    assert any("latin-1 fallback" in warning for warning in result.extraction_warnings)


def test_unsupported_file_type_error(tmp_path: Path) -> None:
    sample = tmp_path / "data.csv"
    sample.write_text("a,b,c", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_document(str(sample))


def test_missing_file_error(tmp_path: Path) -> None:
    missing = tmp_path / "does_not_exist.txt"

    with pytest.raises(FileNotFoundError, match="File not found"):
        extract_document(str(missing))


def test_checksum_is_generated(tmp_path: Path) -> None:
    sample = tmp_path / "checksum.txt"
    content = "checksum payload"
    sample.write_text(content, encoding="utf-8")

    expected = hashlib.sha256(content.encode("utf-8")).hexdigest()

    assert compute_checksum_sha256(str(sample)) == expected

    result = extract_document(str(sample))
    assert result.checksum_sha256 == expected


def test_file_metadata_is_populated(tmp_path: Path) -> None:
    sample = tmp_path / "meta.txt"
    sample.write_text("metadata test", encoding="utf-8")

    result = extract_document(str(sample))

    assert result.file_path == str(sample.resolve())
    assert result.file_name == "meta.txt"
    assert result.file_type == "txt"
    assert result.file_size_bytes == sample.stat().st_size
    assert isinstance(result.metadata, dict)
    assert isinstance(result.extraction_warnings, list)


def test_docx_extraction(tmp_path: Path) -> None:
    sample = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Executive summary paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"
    document.core_properties.title = "Quarterly Report"
    document.save(str(sample))

    result = extract_document(str(sample))

    assert result.file_type == "docx"
    assert "Executive summary paragraph." in result.text
    assert "Metric | Value" in result.text
    assert "Revenue | 100" in result.text
    assert result.page_count is None
    assert result.metadata.get("title") == "Quarterly Report"
